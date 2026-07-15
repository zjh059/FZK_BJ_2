"""提取现有手册的结构、文字、样式和图片，供重制 2.0 使用。"""

from __future__ import annotations

import json
import sys
import zipfile
from pathlib import Path

from docx import Document


def main() -> None:
    source = Path(sys.argv[1]).resolve()
    output_dir = Path(sys.argv[2]).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)

    document = Document(source)
    report: dict[str, object] = {
        "source": str(source),
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "inline_shape_count": len(document.inline_shapes),
        "sections": [],
        "paragraphs": [],
        "tables": [],
        "headers": [],
        "footers": [],
        "styles_in_use": {},
        "core_properties": {
            "title": document.core_properties.title,
            "subject": document.core_properties.subject,
            "author": document.core_properties.author,
            "last_modified_by": document.core_properties.last_modified_by,
            "created": str(document.core_properties.created or ""),
            "modified": str(document.core_properties.modified or ""),
        },
    }

    styles_in_use: dict[str, int] = {}
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.replace("\u00a0", " ").strip()
        style_name = paragraph.style.name if paragraph.style is not None else ""
        styles_in_use[style_name] = styles_in_use.get(style_name, 0) + 1
        if text or paragraph._p.xpath(".//w:drawing | .//w:pict"):
            report["paragraphs"].append(
                {
                    "index": index,
                    "style": style_name,
                    "alignment": str(paragraph.alignment),
                    "text": text,
                    "has_drawing": bool(paragraph._p.xpath(".//w:drawing | .//w:pict")),
                    "runs": [
                        {
                            "text": run.text,
                            "bold": run.bold,
                            "italic": run.italic,
                            "font": run.font.name,
                            "size_pt": run.font.size.pt if run.font.size else None,
                        }
                        for run in paragraph.runs
                        if run.text
                    ],
                }
            )

    report["styles_in_use"] = styles_in_use

    for table_index, table in enumerate(document.tables):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        report["tables"].append(
            {
                "index": table_index,
                "style": table.style.name if table.style is not None else "",
                "rows": len(table.rows),
                "columns": len(table.columns),
                "data": rows,
            }
        )

    for section_index, section in enumerate(document.sections):
        report["sections"].append(
            {
                "index": section_index,
                "page_width_in": section.page_width.inches,
                "page_height_in": section.page_height.inches,
                "top_margin_in": section.top_margin.inches,
                "right_margin_in": section.right_margin.inches,
                "bottom_margin_in": section.bottom_margin.inches,
                "left_margin_in": section.left_margin.inches,
                "header_distance_in": section.header_distance.inches,
                "footer_distance_in": section.footer_distance.inches,
            }
        )
        report["headers"].append([p.text for p in section.header.paragraphs if p.text])
        report["footers"].append([p.text for p in section.footer.paragraphs if p.text])

    report["inline_shapes"] = [
        {
            "index": i,
            "type": str(shape.type),
            "width_in": shape.width.inches,
            "height_in": shape.height.inches,
        }
        for i, shape in enumerate(document.inline_shapes)
    ]

    media_dir = output_dir / "media"
    media_dir.mkdir(exist_ok=True)
    extracted_media: list[dict[str, object]] = []
    with zipfile.ZipFile(source) as archive:
        for member in archive.namelist():
            if not member.startswith("word/media/"):
                continue
            name = Path(member).name
            target = media_dir / name
            target.write_bytes(archive.read(member))
            extracted_media.append({"name": name, "bytes": target.stat().st_size})
    report["media"] = extracted_media

    report_path = output_dir / "inspection.json"
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(report_path)


if __name__ == "__main__":
    main()
